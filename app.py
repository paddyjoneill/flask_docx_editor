from flask import Flask, render_template, redirect, request, send_file
from docx import Document
import csv
import pandas as pd

app = Flask(__name__)

file = ''
document = ''
placeholders = {}
replacements = {
    "company": "Shell",
    "field": "Auk"
}


@app.route('/')
def index():
    string = 'Hello There'
    return render_template('index.html', string=string)

@app.route('/load_file', methods=['POST', 'GET'])
def load_file():
    if request.method == 'POST':
        file = request.files['file']
        data = pd.read_csv(file)
        print(data)
        print(type(data))
        # with open(file, "r") as csvfile:
        #     reader = csv.reader(csvfile, skipinitialspace=True)
        #     for row in reader:
        #         placeholders[row[0]] = row[1]
        #         print(placeholders)
        return render_template('load_file.html', file=file.filename)
    return render_template('load_file.html')

@app.route('/load_template', methods=['POST', 'GET'])
def load_template():
    if request.method == 'POST':
        template = request.files['file']
        document = Document(template)
        filename = request.files['file'].filename
        document_change()
        return render_template('load_template.html', file=filename)
    return render_template('load_template.html')

@app.route('/download')
def download_file():
    path="./new_file.docx"
    return send_file(path, as_attachment=True)


@app.route('/addreplacement', methods=['POST', 'GET'])
def add_replacement():
    if request.method == 'POST':
        replacements[request.form['key']] = request.form['value'] 
    return render_template('addreplacement.html', replacements=replacements)

@app.route('/update/<key>')
def update(key):
    return render_template('addreplacement.html', replacements=replacements, keytoupdate=key)

@app.route('/delete/<key>')
def delete(key):
    del replacements[key]
    return render_template('addreplacement.html', replacements=replacements)



def document_change():
    document = Document('test_file.docx')
    for paragraph in document.paragraphs:
        string = paragraph.text
        for replacement in replacements:
            string = string.replace('#' + replacement, replacements[replacement])
        paragraph.text = string
    document.save('new_file.docx')
    return document

# if __name__ == '__main__':
#     app.run()